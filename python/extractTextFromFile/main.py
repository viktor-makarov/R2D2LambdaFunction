import http.client
import urllib.parse
from io import BytesIO
from docx import Document
import openpyxl
import PyPDF2


def download_file(url):
    # Parse the URL to extract components
    url_parts = urllib.parse.urlparse(url)
    conn = http.client.HTTPSConnection(url_parts.netloc)
    conn.request("GET", url_parts.path)
    response = conn.getresponse()
    
    if response.status != 200:
        raise Exception(f"Failed to download file: {url} {response.status} {response.reason}")

    binaryObject = BytesIO(response.read())
    binaryObject.name = url_parts.path.split('/')[-1]
    
    # Read the response data into a BytesIO object
    return binaryObject

def extracttextfromtxt(txt_stream):
    return txt_stream.read().decode('utf-8')

def extracttextfromxlsx(xlsx_stream):
    wb = openpyxl.load_workbook(xlsx_stream, data_only=True)
    text = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows():
            text.append(' '.join([str(cell.value) if cell.value is not None else '' for cell in row]))
    return "\n".join(text)

def extracttextfromdocx(docx_stream):
    doc = Document(docx_stream)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def extracttextfromxlsx(xlsx_stream):
    wb = openpyxl.load_workbook(xlsx_stream, data_only=True)
    text = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows():
            text.append(' '.join([str(cell.value) if cell.value is not None else '' for cell in row]))
    return "\n".join(text)

def extracttextfrompdf(pdf_stream):

    text = ''
    pdf_reader = PyPDF2.PdfReader(pdf_stream)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extracttextfromtxt(filestream):
    return filestream.read().decode('utf-8')

def extracttextfromdocx(docx_stream):
    doc = Document(docx_stream)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

    doc = Document(docx_stream)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def ocrPDF(filestream):
    return "OCR yе реализован"    

def extracttextfromfile(filestream, mime_type):

    if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        text = extracttextfromdocx(filestream)
    elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
        text = extracttextfromxlsx(filestream)
    elif mime_type.startswith('text/') or mime_type == 'application/json':
        text = extracttextfromtxt(filestream)
    elif mime_type == 'application/pdf':
        text = extracttextfrompdf(filestream)
        noTextFound = len(text) < 10

        if noTextFound:
            text = ocrPDF(filestream)
        
    elif mime_type == 'image/jpeg':
        text = ocrImage(filestream)
    else:
        raise ValueError("Unsupported file type")
    
    return text

def extractTextFromFileRouter(event, context):
  
    url = event["file_url"]
    
    try:
        filestream = download_file(url)
    except Exception as e:
        return {
            'statusCode': 500,
            'body': {"error":'Could not fetch the file by url provided',"success":0}
        }
    mime_type  = event["file_mime_type"]

    try:
        text = extracttextfromfile(filestream, mime_type)
    except Exception as e:
        return {
            'statusCode': 500,
            'body': {"error":e,"success":0}
        }

    return {
        'statusCode': 200,
        'body': {"text":text,"success":1}
    }